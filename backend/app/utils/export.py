"""简历导出工具 - Markdown/结构化数据 → PDF / DOCX"""

import os
import uuid

import markdown as md

from app.config.settings import settings

# ── 通用 CSS ──────────────────────────────────────

RESUME_CSS = """
body {
    font-family: "Noto Sans SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
    max-width: 800px;
    margin: 0 auto;
    padding: 40px;
}
h1 {
    font-size: 22pt;
    color: #1a1a1a;
    border-bottom: 2px solid #333;
    padding-bottom: 8px;
    margin-bottom: 4px;
}
h2 {
    font-size: 14pt;
    color: #2c3e50;
    border-bottom: 1px solid #ddd;
    padding-bottom: 4px;
    margin-top: 16px;
}
h3 { font-size: 12pt; color: #34495e; }
hr { border: none; border-top: 1px solid #ddd; margin: 8px 0; }
ul { padding-left: 20px; }
li { margin-bottom: 4px; }
p { margin: 6px 0; }
strong { color: #1a1a1a; }
"""

# ── 模板 CSS 定义 ──────────────────────────────────────

TEMPLATE_CSS = {
    "professional": """
body { font-family: "Noto Sans SC", "Georgia", serif; font-size: 11pt; color: #333; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 40px; }
.header { background: #1e3a5f; color: white; padding: 24px 32px; margin: -40px -40px 24px; }
.header h1 { font-size: 24pt; margin: 0 0 4px; border: none; color: white; }
.header .title { font-size: 13pt; opacity: 0.9; }
.header .contact { font-size: 10pt; opacity: 0.8; margin-top: 8px; }
.header .contact span { margin-right: 16px; }
.section { margin-bottom: 16px; }
.section h2 { font-size: 13pt; color: #1e3a5f; border-bottom: 2px solid #1e3a5f; padding-bottom: 4px; margin: 0 0 10px; text-transform: uppercase; letter-spacing: 1px; }
.entry { margin-bottom: 12px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 2px; }
.entry-header h3 { font-size: 11pt; color: #1a1a1a; margin: 0; }
.entry-header .date { font-size: 10pt; color: #666; white-space: nowrap; }
.entry-sub { font-size: 10pt; color: #555; margin-bottom: 4px; }
ul { padding-left: 18px; margin: 4px 0; }
li { font-size: 10.5pt; margin-bottom: 3px; line-height: 1.5; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 8px 24px; }
.skill-group { font-size: 10.5pt; }
.skill-group strong { color: #1e3a5f; }
.summary { font-size: 10.5pt; line-height: 1.6; color: #444; font-style: italic; }
""",
    "minimalist": """
body { font-family: "Noto Sans SC", "Helvetica Neue", Arial, sans-serif; font-size: 10.5pt; color: #222; margin: 0; padding: 0; }
.resume { max-width: 720px; margin: 0 auto; padding: 48px 40px; }
.header { text-align: center; margin-bottom: 24px; border-bottom: 1px solid #ddd; padding-bottom: 16px; }
.header h1 { font-size: 28pt; font-weight: 300; letter-spacing: 4px; margin: 0 0 6px; border: none; color: #111; }
.header .title { font-size: 11pt; color: #666; letter-spacing: 2px; }
.header .contact { font-size: 9.5pt; color: #888; margin-top: 10px; }
.header .contact span { margin: 0 8px; }
.section { margin-bottom: 18px; }
.section h2 { font-size: 10pt; color: #999; text-transform: uppercase; letter-spacing: 3px; border-bottom: none; margin: 0 0 10px; font-weight: 400; }
.entry { margin-bottom: 12px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 10.5pt; font-weight: 600; margin: 0; }
.entry-header .date { font-size: 9.5pt; color: #999; }
.entry-sub { font-size: 9.5pt; color: #666; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; color: #444; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 6px 20px; font-size: 10pt; }
.summary { font-size: 10pt; color: #555; line-height: 1.7; text-align: center; }
""",
    "creative": """
body { font-family: "Noto Sans SC", "Helvetica Neue", sans-serif; font-size: 10.5pt; color: #333; margin: 0; padding: 0; }
.resume { display: flex; min-height: 100vh; }
.sidebar { width: 240px; background: linear-gradient(135deg, #6366f1, #8b5cf6); color: white; padding: 32px 20px; flex-shrink: 0; }
.sidebar h1 { font-size: 20pt; margin: 0 0 4px; border: none; color: white; }
.sidebar .title { font-size: 11pt; opacity: 0.85; margin-bottom: 20px; }
.sidebar .contact { font-size: 9.5pt; opacity: 0.8; }
.sidebar .contact div { margin-bottom: 6px; }
.sidebar .section h2 { font-size: 10pt; color: rgba(255,255,255,0.7); text-transform: uppercase; letter-spacing: 2px; border-bottom: 1px solid rgba(255,255,255,0.2); padding-bottom: 4px; margin: 16px 0 8px; }
.sidebar .skill-item { display: flex; justify-content: space-between; font-size: 9.5pt; margin-bottom: 4px; }
.sidebar .skill-bar { height: 4px; background: rgba(255,255,255,0.2); border-radius: 2px; margin-top: 2px; }
.sidebar .skill-bar-fill { height: 100%; background: white; border-radius: 2px; }
.main { flex: 1; padding: 32px; }
.main .section { margin-bottom: 18px; }
.main .section h2 { font-size: 12pt; color: #6366f1; border-bottom: 2px solid #e0e0ff; padding-bottom: 4px; margin: 0 0 10px; }
.entry { margin-bottom: 14px; padding-left: 12px; border-left: 3px solid #e0e0ff; }
.entry-header h3 { font-size: 11pt; margin: 0; color: #1a1a1a; }
.entry-header .date { font-size: 9.5pt; color: #8b5cf6; }
.entry-sub { font-size: 9.5pt; color: #666; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; }
.summary { font-size: 10.5pt; color: #444; line-height: 1.6; }
""",
    "tech": """
body { font-family: "Noto Sans SC", "Fira Code", "JetBrains Mono", monospace; font-size: 10.5pt; color: #e0e0e0; background: #0f172a; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 40px; }
.header { background: #1e293b; border: 1px solid #334155; border-radius: 8px; padding: 24px; margin-bottom: 20px; }
.header h1 { font-size: 22pt; color: #22d3ee; margin: 0 0 4px; border: none; font-family: "Fira Code", monospace; }
.header .title { font-size: 11pt; color: #94a3b8; }
.header .contact { font-size: 9.5pt; color: #64748b; margin-top: 8px; }
.header .contact span { margin-right: 16px; }
.header .contact a { color: #22d3ee; text-decoration: none; }
.section { margin-bottom: 16px; }
.section h2 { font-size: 11pt; color: #22d3ee; border-bottom: 1px solid #334155; padding-bottom: 4px; margin: 0 0 10px; font-family: "Fira Code", monospace; }
.section h2::before { content: "// "; color: #475569; }
.entry { margin-bottom: 12px; background: #1e293b; border-radius: 6px; padding: 12px 16px; border: 1px solid #334155; }
.entry-header h3 { font-size: 10.5pt; color: #f1f5f9; margin: 0; }
.entry-header .date { font-size: 9.5pt; color: #22d3ee; font-family: "Fira Code", monospace; }
.entry-sub { font-size: 9.5pt; color: #94a3b8; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; color: #cbd5e1; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 6px; }
.skill-tag { background: #1e293b; border: 1px solid #22d3ee; color: #22d3ee; padding: 2px 10px; border-radius: 4px; font-size: 9.5pt; font-family: "Fira Code", monospace; }
.summary { font-size: 10pt; color: #94a3b8; line-height: 1.6; }
@media print { body { background: white; color: #333; } .resume { padding: 20px; } .header, .entry { background: #f8f9fa; border-color: #ddd; } .section h2, .header h1, .entry-header .date, .skill-tag { color: #0891b2; } .entry-header h3 { color: #1a1a1a; } li, .summary, .entry-sub { color: #444; } .header .contact { color: #666; } .skill-tag { border-color: #0891b2; } }
""",
    "academic": """
body { font-family: "Noto Serif SC", "Times New Roman", "Georgia", serif; font-size: 11pt; color: #1a1a1a; margin: 0; padding: 0; }
.resume { max-width: 760px; margin: 0 auto; padding: 48px 40px; }
.header { text-align: center; margin-bottom: 20px; }
.header h1 { font-size: 20pt; font-weight: 400; letter-spacing: 2px; margin: 0 0 6px; border: none; }
.header .title { font-size: 11pt; color: #555; }
.header .contact { font-size: 10pt; color: #666; margin-top: 8px; }
.header .contact span { margin: 0 6px; }
.section { margin-bottom: 14px; }
.section h2 { font-size: 11pt; font-variant: small-caps; letter-spacing: 2px; border-bottom: 1px solid #333; padding-bottom: 2px; margin: 0 0 8px; font-weight: 400; }
.entry { margin-bottom: 10px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 11pt; font-weight: 600; margin: 0; font-style: italic; }
.entry-header .date { font-size: 10pt; color: #555; }
.entry-sub { font-size: 10pt; color: #444; margin-left: 16px; }
ul { padding-left: 20px; margin: 4px 0; }
li { font-size: 10.5pt; margin-bottom: 2px; }
.summary { font-size: 10.5pt; line-height: 1.7; text-indent: 2em; }
.skills-grid { font-size: 10.5pt; }
.skill-group { margin-bottom: 4px; }
""",
    "executive": """
body { font-family: "Noto Sans SC", "Garamond", "Georgia", serif; font-size: 11pt; color: #2c2c2c; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 48px 44px; }
.header { border-bottom: 2px solid #b8860b; padding-bottom: 16px; margin-bottom: 24px; }
.header h1 { font-size: 26pt; font-weight: 300; letter-spacing: 3px; text-transform: uppercase; margin: 0 0 4px; border: none; color: #1a1a1a; }
.header .title { font-size: 12pt; color: #b8860b; letter-spacing: 1px; }
.header .contact { font-size: 10pt; color: #777; margin-top: 10px; }
.header .contact span { margin-right: 16px; }
.section { margin-bottom: 20px; }
.section h2 { font-size: 11pt; color: #b8860b; text-transform: uppercase; letter-spacing: 2px; border-bottom: 1px solid #e0d5b8; padding-bottom: 4px; margin: 0 0 12px; font-weight: 400; }
.entry { margin-bottom: 14px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 11.5pt; font-weight: 600; margin: 0; color: #1a1a1a; }
.entry-header .date { font-size: 10pt; color: #888; }
.entry-sub { font-size: 10pt; color: #666; font-style: italic; }
ul { padding-left: 18px; margin: 6px 0; }
li { font-size: 10.5pt; margin-bottom: 4px; line-height: 1.6; }
.summary { font-size: 11pt; line-height: 1.7; color: #333; border-left: 3px solid #b8860b; padding-left: 16px; font-style: italic; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 8px 24px; font-size: 10.5pt; }
.skill-group strong { color: #b8860b; }
""",
}


def _render_summary(data: dict) -> str:
    """渲染个人简介模块"""
    if not data.get("summary"):
        return ""
    return f'<div class="section"><h2>个人简介</h2><p class="summary">{data["summary"]}</p></div>'


def _render_work_experience(data: dict) -> str:
    """渲染工作经历模块"""
    if not data.get("work_experience"):
        return ""
    entries = ""
    for exp in data["work_experience"]:
        hl = "".join(f"<li>{h}</li>" for h in exp.get("highlights", []))
        entries += f"""<div class="entry">
            <div class="entry-header"><h3>{exp['title']}</h3><span class="date">{exp.get('start_date','')} - {exp.get('end_date','')}</span></div>
            <div class="entry-sub">{exp['company']}{(' · ' + exp['location']) if exp.get('location') else ''}</div>
            <ul>{hl}</ul></div>"""
    return f'<div class="section"><h2>工作经历</h2>{entries}</div>'


def _render_education(data: dict) -> str:
    """渲染教育背景模块"""
    if not data.get("education"):
        return ""
    entries = ""
    for edu in data["education"]:
        hl = "".join(f"<li>{h}</li>" for h in edu.get("highlights", []))
        gpa_str = f" | GPA: {edu['gpa']}" if edu.get("gpa") else ""
        entries += f"""<div class="entry">
            <div class="entry-header"><h3>{edu['degree']} - {edu['field']}</h3><span class="date">{edu.get('start_date','')} - {edu.get('end_date','')}</span></div>
            <div class="entry-sub">{edu['institution']}{gpa_str}</div>
            {'<ul>' + hl + '</ul>' if hl else ''}</div>"""
    return f'<div class="section"><h2>教育背景</h2>{entries}</div>'


def _render_skills(data: dict, template_id: str) -> str:
    """渲染专业技能模块"""
    if not data.get("skills"):
        return ""
    if template_id == "tech":
        tags = ""
        for sg in data["skills"]:
            for item in sg.get("items", []):
                tags += f'<span class="skill-tag">{item}</span>'
        return f'<div class="section"><h2>专业技能</h2><div class="skills-grid">{tags}</div></div>'
    else:
        items = ""
        for sg in data["skills"]:
            items += f'<div class="skill-group"><strong>{sg["category"]}</strong>: {", ".join(sg.get("items", []))}</div>'
        return f'<div class="section"><h2>专业技能</h2><div class="skills-grid">{items}</div></div>'


def _render_projects(data: dict, template_id: str) -> str:
    """渲染项目经验模块"""
    if not data.get("projects"):
        return ""
    entries = ""
    for proj in data["projects"]:
        hl = "".join(f"<li>{h}</li>" for h in proj.get("highlights", []))
        tech = ""
        if proj.get("tech_stack"):
            if template_id == "tech":
                tech = '<div style="margin-top:4px">' + "".join(f'<span class="skill-tag">{t}</span>' for t in proj["tech_stack"]) + "</div>"
            else:
                tech = f'<div class="entry-sub">技术栈: {", ".join(proj["tech_stack"])}</div>'
        date_str = ""
        if proj.get("start_date"):
            date_str = f'{proj["start_date"]} - {proj.get("end_date", "")}'
        entries += f"""<div class="entry">
            <div class="entry-header"><h3>{proj['name']}{(' - ' + proj['role']) if proj.get('role') else ''}</h3><span class="date">{date_str}</span></div>
            <p class="entry-sub">{proj.get('description','')}</p>
            {tech}<ul>{hl}</ul></div>"""
    return f'<div class="section"><h2>项目经验</h2>{entries}</div>'


def _render_certifications(data: dict) -> str:
    """渲染证书与奖项模块"""
    if not data.get("certifications"):
        return ""
    items = ""
    for cert in data["certifications"]:
        line = cert["name"]
        if cert.get("issuer"):
            line += f' ({cert["issuer"]})'
        if cert.get("date"):
            line += f' - {cert["date"]}'
        items += f"<li>{line}</li>"
    return f'<div class="section"><h2>证书与奖项</h2><ul>{items}</ul></div>'


def _render_resume_html(data: dict, template_id: str, module_order: list[str] = None) -> str:
    """将结构化数据渲染为 HTML"""
    if module_order is None:
        module_order = ["personal", "summary", "work_experience", "education", "projects", "skills", "certifications"]

    personal = data.get("personal", {})
    is_creative = template_id == "creative"

    # Contact info
    contact_parts = []
    for key in ["phone", "email", "location", "website", "linkedin", "github"]:
        val = personal.get(key)
        if val:
            contact_parts.append(f'<span>{val}</span>')
    contact_html = " ".join(contact_parts)

    # 模块渲染函数映射
    section_renderers = {
        "summary": lambda: _render_summary(data),
        "work_experience": lambda: _render_work_experience(data),
        "education": lambda: _render_education(data),
        "projects": lambda: _render_projects(data, template_id),
        "skills": lambda: _render_skills(data, template_id),
        "certifications": lambda: _render_certifications(data),
    }

    # 按 module_order 顺序渲染
    sections_html = ""
    for module_id in module_order:
        if module_id == "personal":
            continue  # personal 在 header 中处理
        if module_id in section_renderers:
            sections_html += section_renderers[module_id]()

    # Custom Sections
    for sec in data.get("custom_sections", []):
        sections_html += f'<div class="section"><h2>{sec["title"]}</h2><p>{sec["content"]}</p></div>'

    # Assemble based on template layout
    if is_creative:
        # Two-column layout: sidebar + main
        sidebar_sections = ""
        # Put contact and skills in sidebar
        sidebar_contact = "".join(f'<div>{p}</div>' for p in [
            personal.get("phone", ""),
            personal.get("email", ""),
            personal.get("location", ""),
            personal.get("website", ""),
            personal.get("github", ""),
        ] if p)

        sidebar_skills = ""
        if data.get("skills"):
            skill_items = ""
            for sg in data["skills"]:
                for item in sg.get("items", []):
                    skill_items += f'<div class="skill-item"><span>{item}</span></div>'
            sidebar_skills = f'<div class="section"><h2>技能</h2>{skill_items}</div>'

        return f"""<div class="resume">
            <div class="sidebar">
                <h1>{personal.get('name','')}</h1>
                <div class="title">{personal.get('title','')}</div>
                <div class="contact">{sidebar_contact}</div>
                {sidebar_skills}
            </div>
            <div class="main">{sections_html}</div>
        </div>"""
    else:
        return f"""<div class="resume">
            <div class="header">
                <h1>{personal.get('name','')}</h1>
                <div class="title">{personal.get('title','')}</div>
                <div class="contact">{contact_html}</div>
            </div>
            {sections_html}
        </div>"""


# ── PDF 生成器 (fpdf2, 跨平台无需 GTK) ──────────────────────────────

def _make_pdf():
    """创建配置好中文字体的 FPDF 实例, 返回 (pdf, font_name)"""
    from fpdf import FPDF

    pdf = FPDF(unit='pt', format='A4')
    pdf.set_auto_page_break(auto=True, margin=40)

    # 注入中文字体 (macOS / Linux 兼容)
    font_paths = [
        '/System/Library/Fonts/STHeiti Medium.ttc',
        '/System/Library/Fonts/Hiragino Sans GB.ttc',
        '/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc',
        '/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc',
    ]
    chinese_font = None
    for fp in font_paths:
        if os.path.exists(fp):
            chinese_font = fp
            break

    font_name = 'Helvetica'
    if chinese_font:
        try:
            pdf.add_font('chinese', '', chinese_font)
            font_name = 'chinese'
        except Exception:
            pass

    return pdf, font_name


# ── Markdown 导出 (旧版兼容) ──────────────────────────────

def markdown_to_pdf(content: str) -> str:
    """Markdown → PDF, 返回文件路径 (fpdf2)"""
    pdf, font_name = _make_pdf()
    pdf.add_page()

    # 简单 Markdown 渲染
    pdf.set_font(font_name, size=11)
    pdf.set_text_color(51, 51, 51)

    for line in content.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(6)
            continue
        if line.startswith('# '):
            pdf.set_font(font_name, size=18)
            pdf.set_text_color(26, 26, 26)
            pdf.ln(4)
            pdf.multi_cell(0, 20, line[2:], align='L')
            pdf.ln(4)
            pdf.set_font(font_name, size=11)
            pdf.set_text_color(51, 51, 51)
        elif line.startswith('## '):
            pdf.set_font(font_name, size=14)
            pdf.set_text_color(44, 62, 80)
            pdf.ln(6)
            pdf.multi_cell(0, 16, line[3:], align='L')
            pdf.ln(2)
            pdf.set_font(font_name, size=11)
            pdf.set_text_color(51, 51, 51)
        elif line.startswith('- '):
            pdf.set_x(pdf.l_margin + 16)
            pdf.multi_cell(0, 14, '• ' + line[2:], align='L')
        else:
            pdf.multi_cell(0, 14, line, align='L')

    filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    pdf.output(output_path)
    return output_path


def markdown_to_docx(content: str) -> str:
    """Markdown → DOCX, 返回文件路径"""
    import re
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ── 结构化数据导出 (fpdf2) ──────────────────────────────────────

def structured_to_pdf(data: dict, template_id: str = "professional", module_order: list[str] = None) -> str:
    """结构化简历数据 + 模板 → PDF (fpdf2, 跨平台无需 GTK)"""
    if module_order is None:
        module_order = ["summary", "work_experience", "education", "skills", "projects", "certifications"]

    pdf, fn = _make_pdf()
    pdf.add_page()

    # ── 模板样式配置 ──────────────────────────────────
    styles = {
        "professional": dict(
            header_bg=(30, 58, 95), header_fg=(255, 255, 255),
            section_color=(30, 58, 95), entry_main=(26, 26, 26),
            entry_sub=(68, 68, 68), entry_date=(102, 102, 102),
            font=fn, font_size=11,
        ),
        "minimalist": dict(
            header_bg=(255, 255, 255), header_fg=(26, 26, 26),
            section_color=(170, 170, 170), entry_main=(26, 26, 26),
            entry_sub=(119, 119, 119), entry_date=(170, 170, 170),
            font=fn, font_size=10.5,
        ),
        "tech": dict(
            header_bg=(30, 41, 59), header_fg=(34, 211, 238),
            section_color=(34, 211, 238), entry_main=(226, 232, 240),
            entry_sub=(34, 211, 238), entry_date=(100, 116, 139),
            font=fn, font_size=10.5,
        ),
        "academic": dict(
            header_bg=(255, 255, 255), header_fg=(26, 26, 26),
            section_color=(26, 26, 26), entry_main=(26, 26, 26),
            entry_sub=(68, 68, 68), entry_date=(85, 85, 85),
            font=fn, font_size=11,
        ),
        "executive": dict(
            header_bg=(255, 255, 255), header_fg=(26, 26, 26),
            section_color=(184, 134, 11), entry_main=(26, 26, 26),
            entry_sub=(184, 134, 11), entry_date=(136, 136, 136),
            font=fn, font_size=11,
        ),
        "creative": dict(
            header_bg=(99, 102, 241), header_fg=(255, 255, 255),
            section_color=(99, 102, 241), entry_main=(26, 26, 26),
            entry_sub=(99, 102, 241), entry_date=(153, 153, 153),
            font=fn, font_size=10.5,
        ),
    }
    s = styles.get(template_id, styles["professional"])
    W = pdf.w - pdf.l_margin - pdf.r_margin  # 正文字宽

    def set_color(rgb):
        pdf.set_text_color(*rgb)

    def draw_header():
        p = data.get("personal", {})
        name = p.get('name', '')
        title = p.get('title', '')

        # 填背景色
        pdf.set_fill_color(*s['header_bg'])
        pdf.rect(0, 0, pdf.w, 140, 'F')

        set_color(s['header_fg'])
        pdf.set_font(s['font'], size=22)
        pdf.set_xy(pdf.l_margin, 18)
        pdf.cell(W, 28, name, align='C')
        if title:
            pdf.ln(4)
            pdf.set_font(s['font'], size=12)
            pdf.cell(W, 16, title, align='C')

        # 联系信息行
        contact = []
        for k in ['phone', 'email', 'location', 'website', 'linkedin', 'github']:
            v = p.get(k)
            if v:
                contact.append(str(v))
        if contact:
            pdf.ln(4)
            pdf.set_font(s['font'], size=9)
            pdf.cell(W, 14, ' | '.join(contact), align='C')

        pdf.ln(20)

    def draw_section_title(title: str):
        set_color(s['section_color'])
        pdf.set_font(s['font'], 'B', 12)
        pdf.cell(W, 18, title.upper(), align='L')
        pdf.ln(18)
        # 下划线
        y = pdf.get_y()
        pdf.set_draw_color(*s['section_color'])
        pdf.set_line_width(0.8)
        pdf.line(pdf.l_margin, y, pdf.l_margin + W, y)
        pdf.ln(6)

    def draw_entry_header(main: str, date: str = '', x=None):
        y0 = pdf.get_y()
        if x is None:
            x = pdf.l_margin
        fw = W - 120
        set_color(s['entry_main'])
        pdf.set_font(s['font'], 'B', 11)
        pdf.set_x(x)
        pdf.cell(fw, 14, main, align='L')
        if date:
            set_color(s['entry_date'])
            pdf.set_font(s['font'], size=9)
            pdf.cell(120, 14, date, align='R')
        pdf.ln(14)
        set_color(s['entry_sub'])

    def draw_bullets(items: list):
        for item in items:
            set_color(s['entry_sub'])
            pdf.set_x(pdf.l_margin + 14)
            pdf.set_font(s['font'], size=10)
            pdf.multi_cell(W - 14, 12, '• ' + str(item), align='L')

    # 渲染
    draw_header()
    set_color(s['entry_sub'])

    for module_id in module_order:
        if module_id == 'personal':
            continue

        if module_id == 'summary' and data.get('summary'):
            draw_section_title('个人简介')
            set_color(s['entry_sub'])
            pdf.set_font(s['font'], size=s['font_size'])
            pdf.multi_cell(W, 14, str(data['summary']), align='L')
            pdf.ln(8)

        elif module_id == 'work_experience' and data.get('work_experience'):
            draw_section_title('工作经历')
            for exp in data['work_experience']:
                loc = exp.get('location', '')
                date_range = f"{exp.get('start_date','')} - {exp.get('end_date','')}"
                draw_entry_header(
                    f"{exp.get('title','')} @ {exp.get('company','')}{' · '+loc if loc else ''}",
                    date_range
                )
                draw_bullets(exp.get('highlights', []))
                pdf.ln(6)

        elif module_id == 'education' and data.get('education'):
            draw_section_title('教育背景')
            for edu in data['education']:
                gpa = edu.get('gpa', '')
                date_range = f"{edu.get('start_date','')} - {edu.get('end_date','')}"
                detail = f"{edu.get('degree','')} - {edu.get('field','')}"
                if gpa:
                    detail += f' | GPA: {gpa}'
                draw_entry_header(
                    f"{edu.get('institution','')}",
                    date_range
                )
                set_color(s['entry_sub'])
                pdf.set_x(pdf.l_margin)
                pdf.set_font(s['font'], size=s['font_size'])
                pdf.multi_cell(W, 12, detail, align='L')
                draw_bullets(edu.get('highlights', []))
                pdf.ln(4)

        elif module_id == 'skills' and data.get('skills'):
            draw_section_title('专业技能')
            for sg in data['skills']:
                set_color(s['entry_main'])
                pdf.set_font(s['font'], 'B', s['font_size'])
                pdf.cell(80, 14, sg.get('category', ''), align='L')
                set_color(s['entry_sub'])
                pdf.set_font(s['font'], size=s['font_size'])
                pdf.multi_cell(W - 80, 14, ', '.join(sg.get('items', [])), align='L')
            pdf.ln(6)

        elif module_id == 'projects' and data.get('projects'):
            draw_section_title('项目经验')
            for proj in data['projects']:
                role = proj.get('role', '')
                date_range = f"{proj.get('start_date','')} - {proj.get('end_date','')}"
                main = proj.get('name', '')
                if role:
                    main += f' - {role}'
                draw_entry_header(main, date_range)
                if proj.get('description'):
                    set_color(s['entry_sub'])
                    pdf.set_font(s['font'], size=s['font_size'])
                    pdf.multi_cell(W, 12, str(proj['description']), align='L')
                draw_bullets(proj.get('highlights', []))
                tech = proj.get('tech_stack', [])
                if tech:
                    set_color(s['section_color'])
                    pdf.set_font(s['font'], size=9)
                    pdf.set_x(pdf.l_margin)
                    pdf.cell(W, 12, 'Tech: ' + ', '.join(tech), align='L')
                pdf.ln(4)

        elif module_id == 'certifications' and data.get('certifications'):
            draw_section_title('证书与奖项')
            for cert in data['certifications']:
                line = cert.get('name', '')
                detail = []
                if cert.get('issuer'):
                    detail.append(cert['issuer'])
                if cert.get('date'):
                    detail.append(cert['date'])
                set_color(s['entry_main'])
                pdf.set_font(s['font'], 'B', s['font_size'])
                pdf.cell(W, 13, line, align='L')
                if detail:
                    set_color(s['entry_date'])
                    pdf.set_font(s['font'], size=9)
                    pdf.cell(0, 13, ' - '.join(detail), align='L')
                pdf.ln(13)

    # 输出
    filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    pdf.output(output_path)
    return output_path


def structured_to_docx(data: dict, template_id: str = "professional", module_order: list[str] = None) -> str:
    """结构化简历数据 → DOCX"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if module_order is None:
        from app.api.settings_api import _read_config
        import json
        cfg = _read_config()
        raw = cfg.get("module_order", "")
        try:
            module_order = json.loads(raw) if raw else None
        except Exception:
            module_order = None

    if module_order is None:
        module_order = ["personal", "summary", "work_experience", "education", "projects", "skills", "certifications"]

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    personal = data.get("personal", {})

    # Name
    p = doc.add_heading(personal.get("name", ""), level=1)
    if personal.get("title"):
        p = doc.add_paragraph(personal["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact_items = [personal.get(k) for k in ["phone", "email", "location", "website"] if personal.get(k)]
    if contact_items:
        p = doc.add_paragraph(" | ".join(contact_items))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 模块渲染函数映射
    def render_summary():
        if data.get("summary"):
            doc.add_heading("个人简介", level=2)
            doc.add_paragraph(data["summary"])

    def render_work_experience():
        if data.get("work_experience"):
            doc.add_heading("工作经历", level=2)
            for exp in data["work_experience"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{exp['title']}  ")
                run.bold = True
                run.font.size = Pt(11)
                p.add_run(f"@ {exp['company']}").font.size = Pt(10)
                p = doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                for h in exp.get("highlights", []):
                    doc.add_paragraph(h, style="List Bullet")

    def render_education():
        if data.get("education"):
            doc.add_heading("教育背景", level=2)
            for edu in data["education"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{edu['degree']} - {edu['field']}  ")
                run.bold = True
                p.add_run(f"@ {edu['institution']}")
                p = doc.add_paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    def render_skills():
        if data.get("skills"):
            doc.add_heading("专业技能", level=2)
            for sg in data["skills"]:
                p = doc.add_paragraph()
                run = p.add_run(f"{sg['category']}: ")
                run.bold = True
                p.add_run(", ".join(sg.get("items", [])))

    def render_projects():
        if data.get("projects"):
            doc.add_heading("项目经验", level=2)
            for proj in data["projects"]:
                p = doc.add_paragraph()
                run = p.add_run(proj["name"])
                run.bold = True
                if proj.get("description"):
                    doc.add_paragraph(proj["description"])
                for h in proj.get("highlights", []):
                    doc.add_paragraph(h, style="List Bullet")

    def render_certifications():
        if data.get("certifications"):
            doc.add_heading("证书与奖项", level=2)
            for cert in data["certifications"]:
                line = cert["name"]
                if cert.get("issuer"):
                    line += f" ({cert['issuer']})"
                if cert.get("date"):
                    line += f" - {cert['date']}"
                doc.add_paragraph(line, style="List Bullet")

    # 按 module_order 顺序渲染
    section_renderers = {
        "summary": render_summary,
        "work_experience": render_work_experience,
        "education": render_education,
        "projects": render_projects,
        "skills": render_skills,
        "certifications": render_certifications,
    }

    for module_id in module_order:
        if module_id == "personal":
            continue  # personal 已在上面处理
        if module_id in section_renderers:
            section_renderers[module_id]()

    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
