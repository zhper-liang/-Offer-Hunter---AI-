"""文档解析器 - 支持 PDF, DOCX, TXT"""

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """解析后的文档"""
    text: str
    metadata: dict = field(default_factory=dict)


def load_pdf(file_path: str) -> Document:
    """解析 PDF 文件，扫描件自动触发 OCR"""
    from PyPDF2 import PdfReader

    from app.rag.ocr import is_scanned_pdf, try_ocr_for_pdf

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            pages.append(text.strip())

    full_text = "\n\n".join(pages)

    # 如果检测为扫描件（或文字极少），尝试 OCR
    is_ocr = False
    if is_scanned_pdf(file_path, full_text):
        ocr_text = try_ocr_for_pdf(file_path)
        if ocr_text:
            logger.info(f"PDF OCR 成功: {file_path}, 提取文字 {len(ocr_text)} 字符")
            full_text = ocr_text
            is_ocr = True

    return Document(
        text=full_text,
        metadata={
            "filename": Path(file_path).name,
            "file_type": "pdf",
            "page_count": len(reader.pages),
            "is_ocr": is_ocr,
        },
    )


def load_docx(file_path: str) -> Document:
    """解析 DOCX 文件"""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留标题层级信息
            style = para.style.name if para.style else ""
            if style.startswith("Heading"):
                level = style.replace("Heading ", "").strip()
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

    return Document(
        text="\n\n".join(paragraphs),
        metadata={
            "filename": Path(file_path).name,
            "file_type": "docx",
        },
    )


def load_txt(file_path: str) -> Document:
    """解析 TXT 文件"""
    import chardet

    raw = Path(file_path).read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"
    text = raw.decode(encoding, errors="replace")

    return Document(
        text=text.strip(),
        metadata={
            "filename": Path(file_path).name,
            "file_type": "txt",
        },
    )


def load_document(file_path: str) -> Document:
    """根据文件扩展名自动选择解析器"""
    suffix = Path(file_path).suffix.lower()
    loaders = {
        ".pdf": load_pdf,
        ".docx": load_docx,
        ".txt": load_txt,
        ".md": load_txt,
    }
    loader = loaders.get(suffix)
    if not loader:
        raise ValueError(f"不支持的文件类型: {suffix} (支持: {', '.join(loaders)})")
    return loader(file_path)
